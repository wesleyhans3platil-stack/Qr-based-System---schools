#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Word Document Generator for QR Attendance ID Cards
Generates properly formatted Word documents with ID card pairs in a 2x2 grid layout
"""

import sys
import json
import base64
import os
from io import BytesIO
from datetime import datetime
import math

try:
    from docx import Document
    from docx.shared import Inches, Pt, Mm, Cm
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print(json.dumps({"success": False, "error": "python-docx library not installed. Run: pip install python-docx"}))
    sys.exit(1)

def set_cell_border(cell, **kwargs):
    """
    Set cell border with dashed lines for cutting guides
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="dashed" w:sz="4" w:space="0" w:color="999999"/>'
        r'<w:left w:val="dashed" w:sz="4" w:space="0" w:color="999999"/>'
        r'<w:bottom w:val="dashed" w:sz="4" w:space="0" w:color="999999"/>'
        r'<w:right w:val="dashed" w:sz="4" w:space="0" w:color="999999"/>'
        r'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def create_word_document(images_data, output_path, ids_per_page=4):
    """
    Create a Word document with ID card pairs in portrait layout.

    Args:
        images_data: List of base64-encoded PNG images (each image represents an ID pair)
        output_path: Path to save the Word document
        ids_per_page: Number of ID pairs per page (supported: 1,2,4,6)
    """
    # Create document
    doc = Document()
    
    # Set portrait orientation for all sections
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT

    # A4 Portrait dimensions
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    
    # Minimal margins to maximize space
    section.top_margin = Mm(3)
    section.bottom_margin = Mm(3)
    section.left_margin = Mm(3)
    section.right_margin = Mm(3)
    
    # Cell dimensions
    cell_width = Mm(150)
    cell_height = Mm(63)

    # Filter out empty images to avoid creating blank pages
    images_data = [img for img in images_data if img and str(img).strip()]

    # Determine layout based on ids_per_page
    images_per_page = int(ids_per_page) if ids_per_page else 4
    if images_per_page == 1:
        rows_per_page, cols_per_page = 1, 1
    elif images_per_page == 2:
        rows_per_page, cols_per_page = 2, 1
    elif images_per_page == 4:
        rows_per_page, cols_per_page = 2, 2
    elif images_per_page == 6:
        rows_per_page, cols_per_page = 3, 2
    else:
        # Fallback to 4 per page
        rows_per_page, cols_per_page = 2, 2
        images_per_page = rows_per_page * cols_per_page

    total_images = len(images_data)
    total_pages = math.ceil(total_images / images_per_page)

    # Process images page by page; skip pages that would be empty
    for page_idx in range(total_pages):
        start_idx = page_idx * images_per_page
        page_images = images_data[start_idx:start_idx + images_per_page]
        if not page_images:
            continue

        if page_idx > 0:
            doc.add_page_break()

        # Create table for this page (let Word auto-fit to avoid forced page breaks)
        table = doc.add_table(rows=rows_per_page, cols=cols_per_page)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        for r in range(rows_per_page):
            for c in range(cols_per_page):
                img_idx = page_idx * images_per_page + r * cols_per_page + c
                cell = table.rows[r].cells[c]

                # Default empty cell alignment (avoid forcing explicit heights)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                if img_idx >= total_images:
                    continue

                img_data = images_data[img_idx]

                # Extract front and back images
                front_img = None
                back_img = None
                if isinstance(img_data, dict):
                    front_img = img_data.get('front', None)
                    back_img = img_data.get('back', None)
                else:
                    front_img = img_data
                    back_img = img_data

                def _strip_prefix(img):
                    if not img:
                        return None
                    s = str(img)
                    return s.split(',', 1)[1] if ',' in s else s

                front_b64 = _strip_prefix(front_img)
                back_b64 = _strip_prefix(back_img)
                duplicate = (front_b64 and back_b64 and front_b64 == back_b64)

                # Add front image
                front_paragraph = cell.add_paragraph()
                front_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                front_run = front_paragraph.add_run()
                try:
                    if front_b64:
                        front_data = base64.b64decode(front_b64)
                        front_stream = BytesIO(front_data)
                        front_run.add_picture(front_stream, width=Mm(cell_width.mm - 4))
                except Exception as e:
                    front_run.add_text(f"[Front Error: {str(e)[:50]}]")

                # Spacing
                cell.add_paragraph()

                # Add back image only if present and not identical to front
                if back_b64 and not duplicate:
                    back_paragraph = cell.add_paragraph()
                    back_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    back_run = back_paragraph.add_run()
                    try:
                        back_data = base64.b64decode(back_b64)
                        back_stream = BytesIO(back_data)
                        back_run.add_picture(back_stream, width=Mm(cell_width.mm - 4))
                    except Exception as e:
                        back_run.add_text(f"[Back Error: {str(e)[:50]}]")



            # Do not set explicit row heights to avoid layout overflow
            pass
    
    # Save document
    doc.save(output_path)
    return True

def main():
    """
    Main entry point - reads JSON input from stdin or file
    """
    try:
        # Check for input file argument
        if len(sys.argv) > 1:
            input_file = sys.argv[1]
            with open(input_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
        else:
            # Read from stdin
            input_data = sys.stdin.read()
            data = json.loads(input_data)
        
        images = data.get('images', [])
        # Filter out empty image entries to avoid blank pages
        images = [img for img in images if img and str(img).strip()]
        output_path = data.get('output_path', 'output.docx')
        
        if not images:
            print(json.dumps({"success": False, "error": "No images provided"}))
            sys.exit(1)
        
        # Create the document (respect ids_per_page if provided)
        ids_per_page = data.get('ids_per_page', data.get('idsPerPage', 4))
        create_word_document(images, output_path, ids_per_page)

        # Calculate total pages
        pages = math.ceil(len(images) / int(ids_per_page))
        
        print(json.dumps({
            "success": True, 
            "message": f"Document created with {len(images)} ID pairs",
            "output_path": output_path,
            "pages": pages
        }))
        
    except json.JSONDecodeError as e:
        print(json.dumps({"success": False, "error": f"Invalid JSON input: {str(e)}"}))
        sys.exit(1)
    except Exception as e:
        print(json.dumps({"success": False, "error": str(e)}))
        sys.exit(1)

if __name__ == "__main__":
    main()